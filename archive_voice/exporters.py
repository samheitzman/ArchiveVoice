from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

from .constants import QUALITY_WARNING, SPEAKER_WARNING, TRANSLATION_WARNING
from .models import TranscriptResult, TranscriptSegment


POLISH_MARKERS = {
    "ale",
    "babcia",
    "byla",
    "było",
    "były",
    "czekaj",
    "czy",
    "gdzie",
    "jak",
    "jakąś",
    "mi",
    "miały",
    "mnie",
    "nie",
    "od",
    "one",
    "pamięta",
    "pani",
    "przy",
    "ręce",
    "się",
    "tak",
    "tutaj",
    "więźniów",
    "widziała",
    "żydowskich",
}
ENGLISH_MARKERS = {
    "about",
    "any",
    "because",
    "do",
    "does",
    "english",
    "have",
    "here",
    "living",
    "questions",
    "same",
    "say",
    "see",
    "she",
    "starts",
    "the",
    "thing",
    "to",
    "when",
}
POLISH_CHARACTERS = set("ąćęłńóśźżĄĆĘŁŃÓŚŹŻ")
BREAK_AFTER_PHRASES = (
    "dobra, przetłumaczymy",
    "dobra, babcia",
    "dobra babcia",
    "przetłumaczymy",
    "przetłumacz mi",
    "do you have any questions",
    "ok, ok",
)
BREAK_BEFORE_PHRASES = (
    "a jak pani",
    "ale pamięta pani",
    "jak był",
    "jak się",
    "when she",
    "and she",
    "she said",
)


def format_timestamp(seconds: float) -> str:
    total_seconds = max(0, int(seconds))
    hours, remainder = divmod(total_seconds, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def format_duration(seconds: float | None) -> str:
    if seconds is None:
        return "Unknown"
    return format_timestamp(seconds)


def style_slug(style: str) -> str:
    return {
        "Research Transcript": "research_transcript",
        "Timestamped Transcript": "timestamped_transcript",
        "Clean Transcript": "clean_transcript",
        "Reading Transcript": "reading_transcript",
        "Translated Reading": "translated_reading",
    }.get(style, "transcript")


def style_uses_timestamps(style: str, include_timestamps: bool) -> bool:
    if style in {"Research Transcript", "Timestamped Transcript"}:
        return True
    if style in {"Clean Transcript", "Reading Transcript"}:
        return False
    return include_timestamps


def transcript_base_path(audio_path: Path, output_dir: Path, style: str, use_style_suffix: bool = True) -> Path:
    if not use_style_suffix:
        return output_dir / audio_path.stem
    suffix = style_slug(style)
    return output_dir / f"{audio_path.stem}_{suffix}"


def translation_base_path(audio_path: Path, output_dir: Path) -> Path:
    return output_dir / f"{audio_path.stem}_english_translation"


def translated_reading_base_path(audio_path: Path, output_dir: Path) -> Path:
    return output_dir / f"{audio_path.stem}_translated_reading"


def unique_output_path(path: Path) -> Path:
    if not path.exists():
        return path
    for index in range(2, 10_000):
        candidate = path.with_name(f"{path.stem}_{index}{path.suffix}")
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"Could not create a unique output filename for {path.name}.")


def render_transcript_text(result: TranscriptResult, style: str, include_timestamps: bool) -> str:
    metadata = result.metadata
    lines: list[str] = []

    if style == "Research Transcript":
        lines.extend(
            [
                f"Interview: {metadata.source_file}",
                f"Source path: {metadata.source_path}",
                f"Transcribed: {metadata.transcribed_at.isoformat(timespec='seconds')}",
                f"Model: {metadata.model}",
                f"Language mode: {metadata.language_mode}",
                f"Detected language: {metadata.detected_language or 'Not available'}",
                f"Output mode: {metadata.output_mode}",
                f"Duration: {format_duration(metadata.duration_seconds)}",
                "",
            ]
        )

    if style == "Reading Transcript":
        if has_speaker_labels(result.segments):
            lines.extend(render_speaker_reading_paragraphs(result.segments))
        else:
            lines.extend(render_reading_paragraphs(result.segments))
    else:
        for segment in result.segments:
            if style == "Clean Transcript" and not include_timestamps:
                lines.append(format_segment_text(segment))
            elif style == "Research Transcript":
                lines.append(f"[{format_timestamp(segment.start)} - {format_timestamp(segment.end)}]")
                lines.append(format_segment_text(segment))
                lines.append("")
            else:
                lines.append(f"[{format_timestamp(segment.start)}] {format_segment_text(segment)}")

    lines.extend(["", "Notes:", f"- {QUALITY_WARNING}"])
    if has_speaker_labels(result.segments):
        lines.append(f"- {SPEAKER_WARNING}")
    return "\n".join(lines).strip() + "\n"


def render_translation_text(
    result: TranscriptResult,
    include_timestamps: bool,
    source_segments: list[TranscriptSegment] | None = None,
) -> str:
    metadata = result.metadata
    lines = [
        "MACHINE ENGLISH TRANSLATION",
        "This file is a machine-generated English translation.",
        "It is not an original-language transcript and not an English-language transcription.",
        "",
        f"Interview: {metadata.source_file}",
        f"Source path: {metadata.source_path}",
        f"Translated: {metadata.transcribed_at.isoformat(timespec='seconds')}",
        f"Model: {metadata.model}",
        f"Language mode: {metadata.language_mode}",
        f"Detected source language: {metadata.detected_language or 'Not available'}",
        f"Output mode: {metadata.output_mode}",
        f"Duration: {format_duration(metadata.duration_seconds)}",
        "",
        "English translation with source labels:",
        "",
        "Source labels:",
        "- Original English speech: the source transcript at this timestamp appears to be English.",
        "- Machine translation from non-English speech: the source transcript at this timestamp appears to be another language.",
        "- Machine translation, source language uncertain: the source language could not be identified confidently.",
        "",
    ]

    for segment in result.segments:
        source_label = translation_source_label(segment, source_segments)
        speaker_prefix = translation_speaker_prefix(segment, source_segments)
        label = f"{speaker_prefix}{source_label}"
        if include_timestamps:
            lines.append(f"[{format_timestamp(segment.start)} - {format_timestamp(segment.end)}]")
            lines.append(f"{label}: {segment.text}")
            lines.append("")
        else:
            lines.append(f"{label}: {segment.text}")

    lines.extend(["", "Notes:", f"- {TRANSLATION_WARNING}"])
    if has_speaker_labels(source_segments or result.segments):
        lines.append(f"- {SPEAKER_WARNING}")
    return "\n".join(lines).strip() + "\n"


def render_translated_reading_text(
    result: TranscriptResult,
    source_segments: list[TranscriptSegment] | None = None,
) -> str:
    metadata = result.metadata
    lines = [
        "TRANSLATED READING",
        "Machine English translation for reading. This is not an original-language transcript.",
        "",
        f"Interview: {metadata.source_file}",
        f"Translated: {metadata.transcribed_at.isoformat(timespec='seconds')}",
        f"Model: {metadata.model}",
        "",
    ]
    lines.extend(render_translated_reading_paragraphs(result.segments, source_segments))
    lines.extend(["", "Notes:", f"- {TRANSLATION_WARNING}"])
    if has_speaker_labels(source_segments or result.segments):
        lines.append(f"- {SPEAKER_WARNING}")
    return "\n".join(lines).strip() + "\n"


def render_translated_reading_paragraphs(
    segments: list[TranscriptSegment],
    source_segments: list[TranscriptSegment] | None = None,
    pause_break_seconds: float = 1.4,
    max_paragraph_chars: int = 760,
) -> list[str]:
    paragraphs: list[str] = []
    current_parts: list[str] = []
    current_label: str | None = None
    current_group_key: str | None = None
    current_length = 0
    previous_end: float | None = None

    for segment in segments:
        label = translation_block_label(segment, source_segments)
        group_key = translation_block_group_key(segment, source_segments)
        gap = 0.0 if previous_end is None else max(0.0, segment.start - previous_end)
        text = segment.text.strip()
        if not text:
            previous_end = segment.end
            continue
        should_break = bool(current_parts) and (
            group_key != current_group_key
            or gap >= pause_break_seconds
            or current_length + len(text) > max_paragraph_chars
        )
        if should_break:
            append_translated_reading_paragraph(paragraphs, current_label, current_parts)
            current_parts = []
            current_label = None
            current_group_key = None
            current_length = 0
        current_label = current_label or label
        current_group_key = group_key
        current_parts.append(text)
        current_length += len(text) + 1
        previous_end = segment.end

    if current_parts:
        append_translated_reading_paragraph(paragraphs, current_label, current_parts)
    return paragraphs


def append_translated_reading_paragraph(
    paragraphs: list[str],
    label: str | None,
    parts: list[str],
) -> None:
    paragraph = " ".join(parts).strip()
    if not paragraph:
        return
    prefix = f"(Translation) {label}: " if label else "(Translation): "
    paragraphs.append(f"{prefix}{paragraph}")
    paragraphs.append("")


def translation_block_label(
    translated_segment: TranscriptSegment,
    source_segments: list[TranscriptSegment] | None,
) -> str:
    source_label = translation_source_label(translated_segment, source_segments)
    speaker_prefix = translation_speaker_prefix(translated_segment, source_segments)
    return f"{speaker_prefix}{source_label}"


def translation_block_group_key(
    translated_segment: TranscriptSegment,
    source_segments: list[TranscriptSegment] | None,
) -> str:
    speaker_label = translated_segment.speaker_label or speaker_label_from_source(translated_segment, source_segments)
    if speaker_label:
        return speaker_label
    return translation_source_label(translated_segment, source_segments)


def has_speaker_labels(segments: list[TranscriptSegment]) -> bool:
    return any(segment.speaker_label for segment in segments)


def format_segment_text(segment: TranscriptSegment) -> str:
    if segment.speaker_label:
        return f"{segment.speaker_label}: {segment.text}"
    return segment.text


def translation_speaker_prefix(
    translated_segment: TranscriptSegment,
    source_segments: list[TranscriptSegment] | None,
) -> str:
    speaker_label = translated_segment.speaker_label or speaker_label_from_source(translated_segment, source_segments)
    return f"{speaker_label}, " if speaker_label else ""


def speaker_label_from_source(
    translated_segment: TranscriptSegment,
    source_segments: list[TranscriptSegment] | None,
) -> str | None:
    if not source_segments:
        return None
    overlap_by_speaker: dict[str, float] = {}
    for source_segment in overlapping_segments(translated_segment, source_segments):
        if not source_segment.speaker_label:
            continue
        overlap = max(
            0.0,
            min(translated_segment.end, source_segment.end) - max(translated_segment.start, source_segment.start),
        )
        overlap_by_speaker[source_segment.speaker_label] = overlap_by_speaker.get(source_segment.speaker_label, 0.0) + overlap
    if not overlap_by_speaker:
        return None
    return max(overlap_by_speaker.items(), key=lambda item: item[1])[0]


def translation_source_label(
    translated_segment: TranscriptSegment,
    source_segments: list[TranscriptSegment] | None,
) -> str:
    if not source_segments:
        return "Machine translation, source language uncertain"
    source_text = " ".join(segment.text for segment in overlapping_segments(translated_segment, source_segments))
    if not source_text.strip():
        return "Machine translation, source language uncertain"
    language = detect_segment_language(source_text)
    if language == "English":
        return "Original English speech"
    if language is None:
        return "Machine translation, source language uncertain"
    return "Machine translation from non-English speech"


def overlapping_segments(
    segment: TranscriptSegment,
    candidates: list[TranscriptSegment],
    fallback_window_seconds: float = 1.0,
) -> list[TranscriptSegment]:
    overlaps = [
        candidate
        for candidate in candidates
        if max(segment.start, candidate.start) < min(segment.end, candidate.end)
    ]
    if overlaps:
        return overlaps
    return [
        candidate
        for candidate in candidates
        if abs(candidate.start - segment.start) <= fallback_window_seconds
        or abs(candidate.end - segment.end) <= fallback_window_seconds
    ]


def render_reading_paragraphs(
    segments: list[TranscriptSegment],
    pause_break_seconds: float = 1.1,
    max_paragraph_chars: int = 420,
    label_languages: bool = True,
) -> list[str]:
    paragraphs: list[str] = []
    current_parts: list[str] = []
    current_length = 0
    previous_end: float | None = None
    current_language: str | None = None
    units_by_segment = [(segment, split_reading_units(segment.text)) for segment in segments]
    detected_languages = {
        language
        for _, units in units_by_segment
        for language in (detect_segment_language(unit) for unit in units)
        if language is not None
    }
    should_label_languages = label_languages and len(detected_languages) > 1

    for segment, units in units_by_segment:
        for index, text in enumerate(units):
            language = detect_segment_language(text)
            gap = 0.0 if previous_end is None or index > 0 else max(0.0, segment.start - previous_end)
            should_break = bool(current_parts) and (
                gap >= pause_break_seconds
                or language_changed(current_language, language)
                or should_start_new_reading_paragraph(text)
                or should_end_reading_paragraph(current_parts[-1])
                or current_length + len(text) > max_paragraph_chars
            )
            if should_break:
                append_reading_paragraph(paragraphs, current_parts, current_language, should_label_languages)
                current_parts = []
                current_length = 0
                current_language = None
            current_parts.append(text)
            current_length += len(text) + 1
            current_language = merge_language(current_language, language)
        previous_end = segment.end

    if current_parts:
        append_reading_paragraph(paragraphs, current_parts, current_language, should_label_languages)

    return paragraphs


def render_speaker_reading_paragraphs(segments: list[TranscriptSegment]) -> list[str]:
    paragraphs: list[str] = []
    for segment in segments:
        speaker = segment.speaker_label or "Speaker unknown"
        for paragraph_text in render_reading_paragraphs([segment], label_languages=False):
            if paragraph_text:
                paragraphs.append(f"{speaker}: {paragraph_text}")
                paragraphs.append("")
    return paragraphs


def split_reading_units(text: str) -> list[str]:
    stripped = text.strip()
    if not stripped:
        return []
    sentence_units = re.split(r"(?<=[.!?])\s+(?=[A-ZĄĆĘŁŃÓŚŹŻ])", stripped)
    units: list[str] = []
    for sentence in sentence_units:
        units.extend(split_on_reading_cues(sentence))
    return [unit.strip() for unit in units if unit.strip()]


def split_on_reading_cues(text: str) -> list[str]:
    units = [text.strip()]
    for phrase in BREAK_AFTER_PHRASES:
        units = split_after_phrase(units, phrase)
    for phrase in BREAK_BEFORE_PHRASES:
        units = split_before_phrase(units, phrase)
    return units


def split_after_phrase(units: list[str], phrase: str) -> list[str]:
    result: list[str] = []
    pattern = re.compile(rf"(.+?\b{re.escape(phrase)}\b[.!?]?)\s+(.+)", re.IGNORECASE)
    for unit in units:
        match = pattern.match(unit)
        if match:
            result.extend([match.group(1).strip(), match.group(2).strip()])
        else:
            result.append(unit)
    return result


def split_before_phrase(units: list[str], phrase: str) -> list[str]:
    result: list[str] = []
    pattern = re.compile(rf"(.+?)\s+(\b{re.escape(phrase)}\b.+)", re.IGNORECASE)
    for unit in units:
        match = pattern.match(unit)
        if match:
            result.extend([match.group(1).strip(), match.group(2).strip()])
        else:
            result.append(unit)
    return result


def should_start_new_reading_paragraph(text: str) -> bool:
    lowered = text.lower()
    return lowered.endswith("?") or any(lowered.startswith(phrase) for phrase in BREAK_BEFORE_PHRASES)


def should_end_reading_paragraph(text: str) -> bool:
    lowered = text.lower()
    return lowered.endswith("?") or any(phrase in lowered for phrase in BREAK_AFTER_PHRASES)


def append_reading_paragraph(
    paragraphs: list[str],
    parts: list[str],
    language: str | None,
    label_languages: bool,
) -> None:
    paragraph = " ".join(parts).strip()
    if not paragraph:
        return
    if label_languages and language in {"Polish", "English"}:
        paragraph = f"[{language}] {paragraph}"
    paragraphs.append(paragraph)
    paragraphs.append("")


def detect_segment_language(text: str) -> str | None:
    lowered = text.lower()
    words = [word.strip(".,;:!?()[]\"'") for word in lowered.split()]
    polish_score = sum(1 for word in words if word in POLISH_MARKERS)
    english_score = sum(1 for word in words if word in ENGLISH_MARKERS)
    if any(character in POLISH_CHARACTERS for character in text):
        polish_score += 2
    if polish_score >= english_score + 2:
        return "Polish"
    if english_score >= polish_score + 2:
        return "English"
    return None


def language_changed(current_language: str | None, next_language: str | None) -> bool:
    if current_language is None or next_language is None:
        return False
    return current_language != next_language


def merge_language(current_language: str | None, next_language: str | None) -> str | None:
    if current_language is None:
        return next_language
    if next_language is None or current_language == next_language:
        return current_language
    return None


def export_txt(result: TranscriptResult, path: Path, style: str, include_timestamps: bool) -> Path:
    path = unique_output_path(path)
    path.write_text(render_transcript_text(result, style, include_timestamps), encoding="utf-8")
    return path


def export_docx(result: TranscriptResult, path: Path, style: str, include_timestamps: bool) -> Path:
    path = unique_output_path(path)
    try:
        from docx import Document
        from docx.enum.text import WD_BREAK
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Install the app requirements to export DOCX files."
        ) from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    metadata = result.metadata
    document.add_heading(metadata.source_file, level=1)

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Source audio file", metadata.source_file),
        ("Original file path", metadata.source_path),
        ("Date/time transcribed", metadata.transcribed_at.isoformat(timespec="seconds")),
        ("Model used", metadata.model),
        ("Language mode", metadata.language_mode),
        ("Detected language", metadata.detected_language or "Not available"),
        ("Output mode", metadata.output_mode),
        ("Duration", format_duration(metadata.duration_seconds)),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    document.add_paragraph()
    document.add_heading("Transcript", level=2)

    if style == "Reading Transcript":
        paragraphs = (
            render_speaker_reading_paragraphs(result.segments)
            if has_speaker_labels(result.segments)
            else render_reading_paragraphs(result.segments)
        )
        for paragraph_text in paragraphs:
            if paragraph_text:
                document.add_paragraph(paragraph_text)
    elif style == "Clean Transcript" and not include_timestamps:
        for segment in result.segments:
            document.add_paragraph(format_segment_text(segment))
    else:
        for segment in result.segments:
            if style == "Research Transcript":
                stamp = f"[{format_timestamp(segment.start)} - {format_timestamp(segment.end)}]"
                document.add_paragraph(stamp, style=None)
                document.add_paragraph(format_segment_text(segment))
            else:
                document.add_paragraph(f"[{format_timestamp(segment.start)}] {format_segment_text(segment)}")

    document.add_paragraph()
    document.add_heading("Notes", level=2)
    paragraph = document.add_paragraph(QUALITY_WARNING)
    paragraph.runs[-1].add_break(WD_BREAK.LINE)
    if has_speaker_labels(result.segments):
        speaker_paragraph = document.add_paragraph(SPEAKER_WARNING)
        speaker_paragraph.runs[-1].add_break(WD_BREAK.LINE)
    document.add_paragraph("Names, dates, places and unclear passages require human verification.", style=None)

    document.save(path)
    return path


def export_translation_txt(
    result: TranscriptResult,
    path: Path,
    include_timestamps: bool,
    source_segments: list[TranscriptSegment] | None = None,
) -> Path:
    path = unique_output_path(path)
    path.write_text(render_translation_text(result, include_timestamps, source_segments), encoding="utf-8")
    return path


def export_translation_docx(
    result: TranscriptResult,
    path: Path,
    include_timestamps: bool,
    source_segments: list[TranscriptSegment] | None = None,
) -> Path:
    path = unique_output_path(path)
    try:
        from docx import Document
        from docx.enum.text import WD_BREAK
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Install the app requirements to export DOCX files."
        ) from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    metadata = result.metadata
    document.add_heading(f"English Translation - {metadata.source_file}", level=1)
    document.add_paragraph("MACHINE ENGLISH TRANSLATION")
    document.add_paragraph(
        "This file is a machine-generated English translation. It is not an "
        "original-language transcript and not an English-language transcription."
    )

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Source audio file", metadata.source_file),
        ("Original file path", metadata.source_path),
        ("Date/time translated", metadata.transcribed_at.isoformat(timespec="seconds")),
        ("Model used", metadata.model),
        ("Language mode", metadata.language_mode),
        ("Detected source language", metadata.detected_language or "Not available"),
        ("Output mode", metadata.output_mode),
        ("Duration", format_duration(metadata.duration_seconds)),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    document.add_paragraph()
    document.add_heading("English translation with source labels", level=2)
    document.add_paragraph(
        "Source labels distinguish speech that appears to have been English in the "
        "original audio from speech translated from another language."
    )
    for segment in result.segments:
        source_label = translation_source_label(segment, source_segments)
        speaker_prefix = translation_speaker_prefix(segment, source_segments)
        label = f"{speaker_prefix}{source_label}"
        if include_timestamps:
            stamp = f"[{format_timestamp(segment.start)} - {format_timestamp(segment.end)}]"
            document.add_paragraph(stamp, style=None)
        document.add_paragraph(f"{label}: {segment.text}")

    document.add_paragraph()
    document.add_heading("Notes", level=2)
    paragraph = document.add_paragraph(TRANSLATION_WARNING)
    paragraph.runs[-1].add_break(WD_BREAK.LINE)
    if has_speaker_labels(source_segments or result.segments):
        speaker_paragraph = document.add_paragraph(SPEAKER_WARNING)
        speaker_paragraph.runs[-1].add_break(WD_BREAK.LINE)
    document.add_paragraph("Names, dates, places and unclear passages require human verification.", style=None)

    document.save(path)
    return path


def export_translated_reading_txt(
    result: TranscriptResult,
    path: Path,
    source_segments: list[TranscriptSegment] | None = None,
) -> Path:
    path = unique_output_path(path)
    path.write_text(render_translated_reading_text(result, source_segments), encoding="utf-8")
    return path


def export_translated_reading_docx(
    result: TranscriptResult,
    path: Path,
    source_segments: list[TranscriptSegment] | None = None,
) -> Path:
    path = unique_output_path(path)
    try:
        from docx import Document
        from docx.enum.text import WD_BREAK
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Install the app requirements to export DOCX files."
        ) from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    metadata = result.metadata
    document.add_heading(f"Translated Reading - {metadata.source_file}", level=1)
    document.add_paragraph("Machine English translation for reading. This is not an original-language transcript.")

    for paragraph_text in render_translated_reading_paragraphs(result.segments, source_segments):
        if paragraph_text:
            document.add_paragraph(paragraph_text)

    document.add_paragraph()
    document.add_heading("Notes", level=2)
    paragraph = document.add_paragraph(TRANSLATION_WARNING)
    paragraph.runs[-1].add_break(WD_BREAK.LINE)
    if has_speaker_labels(source_segments or result.segments):
        speaker_paragraph = document.add_paragraph(SPEAKER_WARNING)
        speaker_paragraph.runs[-1].add_break(WD_BREAK.LINE)
    document.add_paragraph("Names, dates, places and unclear passages require human verification.", style=None)

    document.save(path)
    return path


def export_json(result: TranscriptResult, path: Path) -> Path:
    path = unique_output_path(path)
    metadata = result.metadata
    payload = {
        "source_file": metadata.source_file,
        "source_path": metadata.source_path,
        "transcribed_at": metadata.transcribed_at.isoformat(timespec="seconds"),
        "model": metadata.model,
        "language_mode": metadata.language_mode,
        "language_code": metadata.language_code,
        "detected_language": metadata.detected_language,
        "detected_language_probability": metadata.detected_language_probability,
        "output_mode": metadata.output_mode,
        "duration_seconds": metadata.duration_seconds,
        "segments": [
            {
                "start": segment.start,
                "end": segment.end,
                "text": segment.text,
                "speaker_label": segment.speaker_label,
            }
            for segment in result.segments
        ],
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def export_all(
    result: TranscriptResult,
    audio_path: Path,
    output_dir: Path,
    styles: list[str],
    include_timestamps: bool,
    write_txt: bool,
    write_docx: bool,
    write_json_sidecar: bool,
) -> list[Path]:
    created: list[Path] = []
    original_styles = [style for style in styles if style != "Translated Reading"]
    use_style_suffix = len(original_styles) > 1
    for style in original_styles:
        base_path = transcript_base_path(audio_path, output_dir, style, use_style_suffix=use_style_suffix)
        timestamps_for_style = style_uses_timestamps(style, include_timestamps)
        if write_txt:
            created.append(export_txt(result, base_path.with_suffix(".txt"), style, timestamps_for_style))
        if write_docx:
            created.append(export_docx(result, base_path.with_suffix(".docx"), style, timestamps_for_style))
    if write_json_sidecar:
        base_path = output_dir / f"{audio_path.stem}_transcript_segments"
        created.append(export_json(result, base_path.with_suffix(".json")))
    return created


def export_translation_all(
    result: TranscriptResult,
    audio_path: Path,
    output_dir: Path,
    include_timestamps: bool,
    write_txt: bool,
    write_docx: bool,
    source_segments: list[TranscriptSegment] | None = None,
    styles: list[str] | None = None,
) -> list[Path]:
    created: list[Path] = []
    translation_styles = styles or ["Detailed Translation"]
    if "Detailed Translation" in translation_styles:
        base_path = translation_base_path(audio_path, output_dir)
        if write_txt:
            created.append(
                export_translation_txt(result, base_path.with_suffix(".txt"), include_timestamps, source_segments)
            )
        if write_docx:
            created.append(
                export_translation_docx(result, base_path.with_suffix(".docx"), include_timestamps, source_segments)
            )
    if "Translated Reading" in translation_styles:
        base_path = translated_reading_base_path(audio_path, output_dir)
        if write_txt:
            created.append(export_translated_reading_txt(result, base_path.with_suffix(".txt"), source_segments))
        if write_docx:
            created.append(export_translated_reading_docx(result, base_path.with_suffix(".docx"), source_segments))
    return created


def empty_result_for_tests(source_file: str = "Interview_01.mp3") -> TranscriptResult:
    from .models import TranscriptMetadata

    return TranscriptResult(
        metadata=TranscriptMetadata(
            source_file=source_file,
            source_path=f"/archive/{source_file}",
            transcribed_at=datetime(2026, 6, 16, 14, 0, 0),
            model="large-v3",
            language_mode="Auto-detect",
            language_code=None,
            detected_language="en",
            duration_seconds=31.0,
        ),
        segments=[
            TranscriptSegment(0.0, 12.0, "My name is Anna."),
            TranscriptSegment(12.0, 31.0, "I remember the winter very clearly."),
        ],
    )
